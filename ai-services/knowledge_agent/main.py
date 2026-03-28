"""
Knowledge Agent Service - RAG-based Knowledge Retrieval

Provides RAG (Retrieval-Augmented Generation) for knowledge base queries.
Migrated from knowledge_assets to knowledge_agent with port 8007.

# Last Update: 2026-03-23 18:40:25
# Author: Daniel Chung
# Version: 2.0.0
"""

import os
from typing import Optional, cast

import httpx
from fastapi import FastAPI, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from pydantic import BaseModel

app = FastAPI(
    title="AIBox Knowledge Agent Service",
    description="RAG-based knowledge retrieval service.",
    version="2.0.0",
)

ALLOWED_ORIGINS = os.getenv(
    "CORS_ORIGINS",
    "http://localhost:1420,http://localhost:6500",
).split(",")

app.add_middleware(
    CORSMiddleware,
    allow_origins=ALLOWED_ORIGINS,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

OLLAMA_BASE_URL = os.getenv("OLLAMA_BASE_URL", "http://localhost:11434")
DEFAULT_MODEL = os.getenv("OLLAMA_MODEL", "qwen2.5-coder:7b")
ARANGO_URL = os.getenv("ARANGO_URL", "http://localhost:8529")
ARANGO_DB = os.getenv("ARANGO_DATABASE", "abc_desktop")
ARANGO_USER = os.getenv("ARANGO_USER", "root")
ARANGO_PASSWORD = os.getenv("ARANGO_PASSWORD", "abc_desktop_2026")


class KnowledgeRequest(BaseModel):
    """Knowledge search request."""

    query: str
    collection: Optional[str] = "knowledge"
    limit: Optional[int] = 5


class Document(BaseModel):
    """A knowledge document."""

    key: str = ""
    content: str
    source: Optional[str] = None
    metadata: Optional[dict[str, object]] = None


class KnowledgeResponse(BaseModel):
    """Knowledge search response."""

    query: str
    results: list[Document]
    context: str
    answer: Optional[str] = None


async def get_embedding(text: str) -> list[float]:
    """Get embedding vector from Ollama."""
    async with httpx.AsyncClient(timeout=30.0) as client:
        response = await client.post(
            f"{OLLAMA_BASE_URL}/api/embed",
            json={"model": "bge-m3:latest", "input": text},
        )
        response.raise_for_status()
        data = response.json()
        embeddings = data.get("embeddings", [])
        if embeddings and len(embeddings) > 0:
            result: list[float] = list(embeddings[0])
            return result
        return []


async def search_similar(collection: str, limit: int) -> list[dict[str, object]]:
    """Search similar documents in ArangoDB collection."""
    aql = f"FOR doc IN {collection} SORT BM25(doc) DESC LIMIT {limit} RETURN doc"
    async with httpx.AsyncClient(timeout=30.0) as client:
        response = await client.post(
            f"{ARANGO_URL}/_db/{ARANGO_DB}/_api/cursor",
            json={"query": aql},
            auth=(ARANGO_USER, ARANGO_PASSWORD),
        )
        response.raise_for_status()
        data = response.json()
        result: list[dict[str, object]] = data.get("result", [])
        return result


async def generate_answer(query: str, context: str) -> str:
    """Generate answer from knowledge context via LLM."""
    prompt = (
        f"Based on the following knowledge base context, "
        f"answer the user's question.\n\n"
        f"Context:\n{context}\n\n"
        f"Question: {query}\n\nAnswer:"
    )
    async with httpx.AsyncClient(timeout=120.0) as client:
        response = await client.post(
            f"{OLLAMA_BASE_URL}/api/chat",
            json={
                "model": DEFAULT_MODEL,
                "messages": [{"role": "user", "content": prompt}],
                "stream": False,
                "temperature": 0.5,
            },
        )
        response.raise_for_status()
        data = response.json()
        content: str = data.get("message", {}).get("content", "").strip()
        return content


@app.get("/")
def root() -> dict[str, str]:
    """Service information."""
    return {
        "service": "knowledge_agent",
        "description": "RAG-based knowledge retrieval",
        "version": "2.0.0",
        "port": "8007",
        "status": "running",
    }


@app.get("/health")
def health() -> dict[str, str]:
    """Health check."""
    return {"status": "ok", "service": "knowledge_agent"}


@app.post("/search", response_model=KnowledgeResponse)
async def search(request: KnowledgeRequest) -> KnowledgeResponse:
    """Search knowledge base and generate answer."""
    try:
        results = await search_similar(
            request.collection or "knowledge",
            request.limit or 5,
        )

        if not results:
            return KnowledgeResponse(
                query=request.query,
                results=[],
                context="",
                answer="No relevant knowledge found.",
            )

        context_parts: list[str] = []
        documents: list[Document] = []

        for doc in results:
            content = str(doc.get("content", ""))
            context_parts.append(content)
            documents.append(
                Document(
                    key=str(doc.get("_key", "")),
                    content=content,
                    source=str(doc.get("source", "")) or None,
                    metadata=None,
                )
            )

        context = "\n\n".join(context_parts)
        answer = await generate_answer(request.query, context)

        return KnowledgeResponse(
            query=request.query,
            results=documents,
            context=context,
            answer=answer,
        )

    except httpx.HTTPError as e:
        raise HTTPException(status_code=502, detail=f"Service unavailable: {str(e)}")
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/add")
async def add_document(
    collection: str, content: str, source: Optional[str] = None
) -> dict[str, object]:
    """Add a document to the knowledge base."""
    doc = {
        "content": content,
        "source": source or "manual",
        "created_at": "2026-03-23",
    }
    async with httpx.AsyncClient(timeout=30.0) as client:
        response = await client.post(
            f"{ARANGO_URL}/_db/{ARANGO_DB}/_api/document/{collection}",
            json=doc,
            auth=(ARANGO_USER, ARANGO_PASSWORD),
        )
        response.raise_for_status()
        result: dict[str, object] = response.json()
        return result


@app.get("/collections")
async def list_collections() -> dict[str, object]:
    """List available collections."""
    async with httpx.AsyncClient(timeout=30.0) as client:
        response = await client.get(
            f"{ARANGO_URL}/_db/{ARANGO_DB}/_api/collection",
            auth=(ARANGO_USER, ARANGO_PASSWORD),
        )
        response.raise_for_status()
        result: dict[str, object] = response.json()
        return result


class PipelineTriggerRequest(BaseModel):
    task: str
    file_id: str
    local_path: str
    root_id: str


@app.post("/pipeline/vector")
async def trigger_vector(file_id: str, root_id: str) -> dict[str, object]:
    from celery_app.tasks import vectorize_task
    from kb_pipeline.arango_ops import ArangoOps

    arango = ArangoOps()
    file_doc = arango.get_file(file_id)
    local_path = file_doc.get("local_path") if file_doc else None
    if not local_path:
        return {"error": "file not found or local_path missing"}
    result = vectorize_task.delay(file_id, local_path, root_id)
    arango.set_task_id(file_id, vector_task_id=result.id)
    return {
        "status": "queued",
        "file_id": file_id,
        "type": "vectorize",
        "task_id": result.id,
    }


class TriggerRequest(BaseModel):
    task: str
    file_id: str
    local_path: str
    root_id: str
    session_key: str | None = None


@app.post("/pipeline/trigger")
async def trigger_pipeline(body: TriggerRequest) -> dict[str, object]:
    from celery_app.tasks import graph_task, vectorize_task
    from kb_pipeline.arango_ops import ArangoOps

    arango = ArangoOps()
    vector_result = vectorize_task.delay(
        body.file_id, body.local_path, body.root_id, session_key=body.session_key
    )
    graph_result = graph_task.delay(
        body.file_id, body.local_path, session_key=body.session_key
    )
    arango.set_task_id(
        body.file_id, vector_task_id=vector_result.id, graph_task_id=graph_result.id
    )
    return {
        "status": "queued",
        "file_id": body.file_id,
        "vector_task_id": vector_result.id,
        "graph_task_id": graph_result.id,
    }


@app.post("/pipeline/graph")
async def trigger_graph(file_id: str) -> dict[str, object]:
    from celery_app.tasks import graph_task
    from kb_pipeline.arango_ops import ArangoOps

    arango = ArangoOps()
    file_doc = arango.get_file(file_id)
    local_path = file_doc.get("local_path") if file_doc else None
    if not local_path:
        return {"error": "file not found or local_path missing"}
    result = graph_task.delay(file_id, local_path)
    arango.set_task_id(file_id, graph_task_id=result.id)
    return {
        "status": "queued",
        "file_id": file_id,
        "type": "graph",
        "task_id": result.id,
    }


@app.get("/pipeline/vectors")
async def get_vectors(
    file_id: str, limit: int = 50, offset: int = 0
) -> dict[str, object]:
    from kb_pipeline.arango_ops import ArangoOps
    from kb_pipeline.qdrant_ops import QdrantStore

    arango = ArangoOps()
    file_doc = arango.get_file(file_id)
    if not file_doc:
        raise HTTPException(status_code=404, detail="file not found")
    root_id = str(file_doc.get("knowledge_root_id", ""))
    if not root_id:
        return {"chunks": [], "total": 0, "file_id": file_id}
    qdrant = QdrantStore()
    collection = f"knowledge_{root_id}"
    chunks = qdrant.get_chunks(collection, file_id, limit, offset)
    return {"chunks": chunks, "total": len(chunks), "file_id": file_id}


@app.get("/pipeline/similar")
async def get_similar(
    file_id: str, chunk_id: str, top_k: int = 10
) -> dict[str, object]:
    from kb_pipeline.arango_ops import ArangoOps
    from kb_pipeline.qdrant_ops import QdrantStore

    arango = ArangoOps()
    file_doc = arango.get_file(file_id)
    if not file_doc:
        raise HTTPException(status_code=404, detail="file not found")
    root_id = str(file_doc.get("knowledge_root_id", ""))
    if not root_id:
        return {"similar": []}

    qdrant = QdrantStore()
    collection = f"knowledge_{root_id}"
    try:
        positive_id = int(chunk_id)
    except ValueError:
        raise HTTPException(status_code=400, detail="chunk_id must be numeric")
    results = qdrant.recommend(collection, positive_id, limit=top_k)
    similar = []
    for r in results:
        pld = cast(dict[str, object], r.get("payload") or {})
        similar.append(
            {
                "chunk_id": str(r.get("id", "")),
                "text": str(pld.get("text_full") or pld.get("text") or ""),
                "score": cast(float, r.get("score", 0.0)),
            }
        )
    return {"similar": similar}


@app.post("/pipeline/regenerate/{file_id}")
async def regenerate_pipeline(file_id: str) -> dict[str, object]:
    from celery_app.tasks import graph_task, vectorize_task
    from kb_pipeline.arango_ops import ArangoOps

    arango = ArangoOps()
    file_doc = arango.get_file(file_id)
    if not file_doc:
        raise HTTPException(status_code=404, detail="file not found")
    local_path = file_doc.get("local_path")
    root_id = file_doc.get("knowledge_root_id")
    if not local_path or not root_id:
        raise HTTPException(
            status_code=400,
            detail="file missing local_path or knowledge_root_id",
        )
    vector_result = vectorize_task.delay(file_id, local_path, root_id)
    graph_result = graph_task.delay(file_id, local_path)
    arango.set_task_id(
        file_id, vector_task_id=vector_result.id, graph_task_id=graph_result.id
    )
    arango.update_status(file_id, vector_status="queued", graph_status="queued")
    return {
        "status": "queued",
        "file_id": file_id,
        "vector_task_id": vector_result.id,
        "graph_task_id": graph_result.id,
    }


@app.get("/pipeline/graph")
async def get_graph(file_id: str) -> dict[str, object]:
    from kb_pipeline.arango_ops import ArangoOps

    arango = ArangoOps()
    file_doc = arango.get_file(file_id)
    if not file_doc:
        raise HTTPException(status_code=404, detail="file not found")
    graph_data = arango.get_graph(file_id)
    return {
        "nodes": graph_data["nodes"],
        "edges": graph_data["edges"],
    }


@app.post("/pipeline/retry")
async def retry_pipeline(file_id: str) -> dict[str, object]:
    from celery_app.tasks import graph_task, vectorize_task
    from kb_pipeline.arango_ops import ArangoOps

    arango = ArangoOps()
    file_doc = arango.get_file(file_id)
    if not file_doc:
        raise HTTPException(status_code=404, detail="file not found")

    local_path = file_doc.get("local_path", "")
    root_id = file_doc.get("knowledge_root_id", "")
    if not local_path:
        raise HTTPException(status_code=400, detail="local_path missing")

    v_result = vectorize_task.delay(file_id, local_path, root_id)
    g_result = graph_task.delay(file_id, local_path)
    arango.set_task_id(file_id, vector_task_id=v_result.id, graph_task_id=g_result.id)
    arango.update_status(file_id, vector_status="pending", graph_status="pending")
    return {
        "status": "queued",
        "file_id": file_id,
        "vector_task_id": v_result.id,
        "graph_task_id": g_result.id,
    }


@app.post("/pipeline/abort")
async def abort_pipeline(file_id: str) -> dict[str, object]:
    from celery_app.app import app as celery_app
    from kb_pipeline.arango_ops import ArangoOps

    arango = ArangoOps()
    file_doc = arango.get_file(file_id)
    if not file_doc:
        raise HTTPException(status_code=404, detail="file not found")

    revoked: list[str] = []
    vector_task_id = file_doc.get("vector_task_id")
    graph_task_id = file_doc.get("graph_task_id")

    for tid in [vector_task_id, graph_task_id]:
        if tid and isinstance(tid, str):
            celery_app.control.revoke(tid, terminate=True)
            revoked.append(tid)

    arango.update_status(
        file_id,
        vector_status="aborted",
        graph_status="aborted",
        failed_reason="任務已被使用者中止",
    )
    return {"status": "aborted", "file_id": file_id, "revoked": revoked}


@app.post("/pipeline/delete")
async def delete_file_data(file_id: str) -> dict[str, object]:
    from pathlib import Path

    from kb_pipeline.arango_ops import ArangoOps
    from kb_pipeline.qdrant_ops import QdrantStore

    arango = ArangoOps()
    file_doc = arango.get_file(file_id)
    if not file_doc:
        raise HTTPException(status_code=404, detail="file not found")

    revoked: list[str] = []
    try:
        import json as json_mod

        import redis as redis_lib

        from celery_app.app import REDIS_URL
        from celery_app.app import app as celery_app

        for tid_key in ("vector_task_id", "graph_task_id"):
            tid = file_doc.get(tid_key)
            if tid and isinstance(tid, str):
                celery_app.control.revoke(tid, terminate=True)
                revoked.append(tid)

        r = redis_lib.from_url(REDIS_URL)
        queue_key = "celery"
        queue_len = r.llen(queue_key)
        if queue_len and queue_len > 0:
            to_remove: list[bytes] = []
            for raw in r.lrange(queue_key, 0, queue_len - 1) or []:
                try:
                    msg = json_mod.loads(raw)
                    body = msg.get("body")
                    if isinstance(body, str):
                        import base64

                        body = json_mod.loads(base64.b64decode(body))
                    args = (
                        body if isinstance(body, list) else (body or {}).get("args", [])
                    )
                    if (
                        isinstance(args, (list, tuple))
                        and len(args) > 0
                        and args[0] == file_id
                    ):
                        task_id = msg.get("headers", {}).get("id", "")
                        if task_id:
                            celery_app.control.revoke(task_id, terminate=True)
                            revoked.append(task_id)
                        to_remove.append(
                            raw if isinstance(raw, bytes) else raw.encode()
                        )
                except Exception:
                    continue
            for item in to_remove:
                r.lrem(queue_key, 1, item)
    except Exception:
        pass

    root_id = str(file_doc.get("knowledge_root_id", ""))
    qdrant_deleted = False
    if root_id:
        try:
            qdrant = QdrantStore()
            qdrant.delete_by_file(f"knowledge_{root_id}", file_id)
            qdrant_deleted = True
        except Exception:
            pass

    arango_removed = arango.delete_file_data(file_id)

    local_deleted = False
    local_path = str(file_doc.get("local_path", ""))
    if local_path and Path(local_path).exists():
        try:
            Path(local_path).unlink()
            local_deleted = True
        except Exception:
            pass

    seaweed_deleted = False
    s3_path = str(file_doc.get("s3_path", ""))
    if s3_path:
        try:
            seaweed_base = os.getenv("SEAWEED_AIBOX_URL", "http://localhost:8888")
            seaweed_user = os.getenv("SEAWEED_USER", "admin")
            seaweed_pass = os.getenv("SEAWEED_PASS", "admin123")
            async with httpx.AsyncClient(timeout=15.0) as client:
                resp = await client.request(
                    "DELETE",
                    f"{seaweed_base}/{s3_path}",
                    auth=(seaweed_user, seaweed_pass),
                )
                seaweed_deleted = resp.status_code < 400
        except Exception:
            pass

    return {
        "status": "deleted",
        "file_id": file_id,
        "revoked_tasks": revoked,
        "qdrant_deleted": qdrant_deleted,
        "arango_removed": arango_removed,
        "local_deleted": local_deleted,
        "seaweed_deleted": seaweed_deleted,
    }


@app.get("/pipeline/logs")
async def get_pipeline_logs(file_id: str) -> dict[str, object]:
    from kb_pipeline.arango_ops import ArangoOps

    arango = ArangoOps()
    logs = arango.get_job_logs(file_id)
    return {"file_id": file_id, "logs": logs, "count": len(logs)}


@app.get("/pipeline/preview")
async def get_preview(file_id: str) -> dict[str, object]:
    from pathlib import Path

    from kb_pipeline.arango_ops import ArangoOps

    arango = ArangoOps()
    file_doc = arango.get_file(file_id)
    if not file_doc:
        raise HTTPException(status_code=404, detail="file not found")

    local_path = str(file_doc.get("local_path"))
    if not local_path or not Path(local_path).exists():
        raise HTTPException(status_code=404, detail="file not found on disk")

    ext = Path(local_path).suffix.lower()
    content_type_map = {
        ".md": "markdown",
        ".txt": "text",
        ".pdf": "text",
        ".csv": "table",
        ".xlsx": "table",
        ".xls": "table",
        ".docx": "text",
        ".doc": "text",
    }
    preview_type = content_type_map.get(ext, "binary")

    if preview_type == "markdown":
        text = Path(local_path).read_text(encoding="utf-8", errors="replace")
        return {"file_id": file_id, "type": "markdown", "content": text}

    if ext == ".pdf":
        download_url = f"/pipeline/download?file_id={file_id}"
        return {"file_id": file_id, "type": "pdf_url", "url": download_url}

    if preview_type == "text":
        text = Path(local_path).read_text(encoding="utf-8", errors="replace")
        return {"file_id": file_id, "type": "text", "content": text[:5000]}

    if preview_type == "table":
        rows: list[dict[str, str | int | float]] = []
        headers: list[str] = []
        if ext in (".xlsx", ".xls"):
            import openpyxl

            wb = openpyxl.load_workbook(local_path, data_only=True)
            ws = wb.active
            for i, row in enumerate(ws.iter_rows(values_only=True)):
                if i == 0:
                    headers = [str(c) if c is not None else "" for c in row]
                else:
                    rows.append(
                        {
                            str(headers[j]) if j < len(headers) else f"col{j}": str(c)
                            if c is not None
                            else ""
                            for j, c in enumerate(row)
                        }
                    )
            return {
                "file_id": file_id,
                "type": "table",
                "headers": headers,
                "rows": rows[:200],
            }
        if ext == ".csv":
            import csv

            with open(local_path, encoding="utf-8", errors="replace") as f:
                reader = csv.DictReader(f)
                for i, row in enumerate(reader):
                    if i == 0:
                        headers = list(row.keys())
                    rows.append(row)
                    if i >= 199:
                        break
            return {
                "file_id": file_id,
                "type": "table",
                "headers": headers,
                "rows": rows,
            }

    if ext == ".pdf":
        download_url = f"/pipeline/download?file_id={file_id}"
        return {"file_id": file_id, "type": "pdf_url", "url": download_url}

    if ext in (".docx", ".doc"):
        from docx import Document

        doc = Document(local_path)
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        return {
            "file_id": file_id,
            "type": "text",
            "content": "\n".join(paragraphs[:200]),
        }

    return {"file_id": file_id, "type": "binary", "message": "不支援的檔案格式"}


@app.get("/pipeline/download")
async def download_file(file_id: str) -> dict[str, object]:
    from pathlib import Path

    from kb_pipeline.arango_ops import ArangoOps

    arango = ArangoOps()
    file_doc = arango.get_file(file_id)
    if not file_doc:
        raise HTTPException(status_code=404, detail="file not found")

    local_path: str = str(file_doc.get("local_path"))
    if not local_path or not Path(local_path).exists():
        raise HTTPException(status_code=404, detail="file not found on disk")

    filename: str = str(file_doc.get("filename", "download"))
    file_type: str = str(file_doc.get("file_type", "application/octet-stream"))
    import mimetypes

    guessed = mimetypes.guess_type(filename)
    mime: str = guessed[0] if guessed[0] else str(file_type)

    return FileResponse(  # type: ignore[return-value]
        path=local_path,
        filename=filename,
        media_type=mime,
        content_disposition_type="inline",
    )
